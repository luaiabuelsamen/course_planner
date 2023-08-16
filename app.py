from flask import Flask, render_template
import argparse
from docx import Document
from consume_cirrculum import parse_courses, extract_table_contents

app = Flask(__name__)

@app.route('/')
def index():
    max_term = max(course['term'] for course in courses) if courses else 0
    return render_template('index.html', courses=courses[:-18], max_term=max_term, program= program)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extract and parse course information from a curriculum table.")
    parser.add_argument("file_path", help="Path to the Word file containing the curriculum table.")
    args = parser.parse_args()

    doc = Document(args.file_path)
    tables = doc.tables
    if doc.paragraphs:
        first_line = doc.paragraphs[0].text
        program = first_line
    else:
        program =  None  # Document has no paragraphs

    if len(tables) > 0:
        all_table_data = []
        for table in tables:
            table_data = extract_table_contents(table)
            all_table_data.extend(table_data)

        courses = parse_courses(all_table_data)
    else:
        courses = []

    app.run()
